"""
报告生成模块，负责将摘要转换为最终的报告文件
"""
from pathlib import Path
from typing import Optional, Union
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..utils.config import config, logger

class ReportGenerator:
    """报告生成类，将摘要转换为格式化的Word文档"""
    
    def __init__(self, 
                 input_file: Optional[Union[str, Path]] = None,
                 output_file: Optional[Union[str, Path]] = None):
        """
        初始化报告生成器
        
        Args:
            input_file: 输入的摘要文件路径，默认使用配置中的路径
            output_file: 输出的报告文件路径，默认使用配置中的路径
        """
        self.input_file = Path(input_file) if input_file else config.summaries_file
        self.output_file = Path(output_file) if output_file else config.report_file
        
    def parse_summaries(self) -> dict:
        """
        解析摘要文件，将摘要内容提取为结构化数据
        
        Returns:
            包含摘要各部分内容的字典
        """
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 合并所有摘要内容
            summaries = []
            for summary_block in re.findall(r'摘要 for Rows.*?(?=摘要 for Rows|\Z)', content, re.DOTALL):
                # 移除标题行
                summary_text = re.sub(r'^摘要 for Rows.*?\n', '', summary_block, flags=re.DOTALL).strip()
                if summary_text:
                    summaries.append(summary_text)
            
            combined_summary = "\n\n".join(summaries)
            
            # 提取各部分内容
            big_things = re.search(r'Big things happening with large language models:(.*?)(?=Some practical tips：|Other things：|\Z)', 
                                  combined_summary, re.DOTALL)
            practical_tips = re.search(r'Some practical tips：(.*?)(?=Other things：|\Z)', 
                                     combined_summary, re.DOTALL)
            other_things = re.search(r'Other things：(.*?)(?=\Z)', 
                                   combined_summary, re.DOTALL)
            
            result = {
                'big_things': big_things.group(1).strip() if big_things else "",
                'practical_tips': practical_tips.group(1).strip() if practical_tips else "",
                'other_things': other_things.group(1).strip() if other_things else "",
            }
            
            return result
            
        except Exception as e:
            logger.error(f"解析摘要文件时出错: {e}")
            return {'big_things': "", 'practical_tips': "", 'other_things': ""}
    
    def generate_docx_report(self) -> bool:
        """
        生成Word格式的报告文档
        
        Returns:
            是否成功生成报告
        """
        logger.info(f"开始生成报告，从文件 {self.input_file} 读取摘要...")
        
        try:
            # 解析摘要内容
            summary_data = self.parse_summaries()
            
            if not any(summary_data.values()):
                logger.error("没有找到有效的摘要内容")
                return False
            
            # 创建Word文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # 添加标题
            title = doc.add_heading(f"{config.current_date} LLM 新闻日报", level=0)
            title_format = title.runs[0].font
            title_format.size = Pt(24)
            title_format.color.rgb = RGBColor(0, 0, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加大型语言模型重要进展部分
            if summary_data['big_things']:
                heading = doc.add_heading("大型语言模型重要进展", level=1)
                heading.runs[0].font.size = Pt(18)
                p = doc.add_paragraph(summary_data['big_things'])
                p.paragraph_format.space_after = Pt(12)
            
            # 添加实用技巧部分
            if summary_data['practical_tips']:
                heading = doc.add_heading("实用技巧", level=1)
                heading.runs[0].font.size = Pt(18)
                p = doc.add_paragraph(summary_data['practical_tips'])
                p.paragraph_format.space_after = Pt(12)
            
            # 添加其他信息部分
            if summary_data['other_things']:
                heading = doc.add_heading("其他信息", level=1)
                heading.runs[0].font.size = Pt(18)
                p = doc.add_paragraph(summary_data['other_things'])
                p.paragraph_format.space_after = Pt(12)
            
            # 添加页脚
            footer = doc.sections[0].footer
            p = footer.paragraphs[0]
            p.text = f"自动生成于 {config.current_date} | LLM 新闻日报"
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 保存文档
            doc.save(self.output_file)
            logger.info(f"报告已成功生成并保存到 {self.output_file}")
            
            return True
            
        except Exception as e:
            logger.error(f"生成报告时出错: {e}")
            return False


def run(input_file: Optional[str] = None, output_file: Optional[str] = None) -> bool:
    """
    运行报告生成的主函数
    
    Args:
        input_file: 可选的输入文件路径
        output_file: 可选的输出文件路径
        
    Returns:
        是否成功生成报告
    """
    try:
        generator = ReportGenerator(input_file, output_file)
        return generator.generate_docx_report()
    except Exception as e:
        logger.error(f"运行报告生成器时出错: {e}")
        return False


if __name__ == "__main__":
    run()